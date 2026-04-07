from __future__ import annotations

import re
from io import BytesIO
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _ensure_analysis(analysis: Any) -> Dict[str, Any]:
    return analysis if isinstance(analysis, dict) else {}


def _replace_raw_once(text: str, raw: str, replacement: str) -> str:
    if not raw or not replacement:
        return text
    return re.sub(re.escape(raw), replacement, text, count=1, flags=re.I)


def _replace_context_once(text: str, old_context: str, new_context: str) -> str:
    if not old_context or not new_context:
        return text
    if old_context in text:
        return text.replace(old_context, new_context, 1)
    return text


def build_marked_text(original_text: str, analysis: Dict[str, Any]) -> str:
    analysis = _ensure_analysis(analysis)
    updated = original_text or ''
    for item in analysis.get('citation_results', []):
        raw = item.get('raw')
        replacement = item.get('substituicao_textual')
        rewritten = item.get('paragrafo_reescrito')
        context = item.get('contexto')
        if rewritten and context and item.get('status') in {'divergente', 'valida_pouco_compativel'}:
            marked = f"[[PARÁGRAFO REESCRITO]] {rewritten}"
            updated = _replace_context_once(updated, context, marked)
        elif replacement and raw and item.get('status') in {'divergente', 'valida_pouco_compativel'}:
            marked = f"[[CORRIGIDO: {replacement}]]"
            updated = _replace_raw_once(updated, raw, marked)
    return updated


def build_revised_text(original_text: str, analysis: Dict[str, Any], mode: str = 'premium') -> str:
    analysis = _ensure_analysis(analysis)
    updated = original_text or ''
    for item in analysis.get('citation_results', []):
        raw = item.get('raw')
        replacement = item.get('substituicao_textual')
        rewritten = item.get('paragrafo_reescrito')
        context = item.get('contexto')
        if item.get('status') not in {'divergente', 'valida_pouco_compativel'}:
            continue
        if mode in {'contextual', 'premium'} and rewritten and context:
            updated = _replace_context_once(updated, context, rewritten)
        elif replacement and raw:
            updated = _replace_raw_once(updated, raw, replacement)
    return updated


def _resolve_title_and_analysis(arg2: Any = None, arg3: Any = None) -> tuple[Dict[str, Any], str]:
    analysis: Dict[str, Any] = {}
    title = 'Peça revisada'
    if isinstance(arg2, dict):
        analysis = arg2
        if isinstance(arg3, str) and arg3.strip():
            title = arg3
    elif isinstance(arg2, str):
        title = arg2 or title
        if isinstance(arg3, dict):
            analysis = arg3
    elif isinstance(arg3, dict):
        analysis = arg3
    return analysis, title


def build_docx_bytes(revised_text: str, arg2: Any = None, arg3: Any = None, marked: bool = False) -> bytes:
    analysis, title = _resolve_title_and_analysis(arg2, arg3)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    doc.add_heading(title, level=1)
    info = doc.add_paragraph()
    info.add_run('Tipo identificado: ').bold = True
    piece_type = analysis.get('piece_type', {}) if isinstance(analysis, dict) else {}
    info.add_run((piece_type or {}).get('tipo', 'Não identificado'))
    info = doc.add_paragraph()
    info.add_run('Foco da versão: ').bold = True
    info.add_run('auditoria de citação, correção automática e reescrita jurídica contextual.')
    doc.add_paragraph('')
    for part in (revised_text or '').split('\n'):
        if not part.strip():
            continue
        p = doc.add_paragraph()
        run = p.add_run(part.strip())
        if marked and ('[[CORRIGIDO:' in part or '[[PARÁGRAFO REESCRITO]]' in part):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_pdf_bytes(revised_text: str, arg2: Any = None, arg3: Any = None) -> bytes:
    analysis, title = _resolve_title_and_analysis(arg2, arg3)
    bio = BytesIO()
    pdf = SimpleDocTemplate(bio, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    story = [Paragraph(f'<b>{title}</b>', styles['Title']), Spacer(1, 12)]
    piece_type = analysis.get('piece_type', {}) if isinstance(analysis, dict) else {}
    story.append(Paragraph(f"<b>Tipo identificado:</b> {(piece_type or {}).get('tipo', 'Não identificado')}", styles['Normal']))
    story.append(Spacer(1, 10))
    safe_text = (revised_text or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    for para in safe_text.split('\n'):
        if para.strip():
            story.append(Paragraph(para.strip(), styles['Normal']))
            story.append(Spacer(1, 6))
    pdf.build(story)
    return bio.getvalue()
